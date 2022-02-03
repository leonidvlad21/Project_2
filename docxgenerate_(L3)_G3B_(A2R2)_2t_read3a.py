#!/usr/bin/python3

# void blank template

import csv
from docx import Document
import copy
import re

from excel_mod_A2R2_3 import get_data

###
# input csv parameters
###

csv_file_name = 'table.csv'
csv_dialect = csv.excel
csv_dialect.doublequote = True
csv_dialect.delimiter = ";"

readings_fname = 'readings_log_A2R2_t2_a.xlsx'

###
# template file
###

template_file_name = 'template_G3B144_A2R2_2tar_read1.docx'

###
# output folder
###

out_folder = 'output'

import os
if not os.path.exists(out_folder):
    os.makedirs(out_folder)

###
# functions that generate field values
###

fld_funcs = {}

###
# fld_funcs["accur"] = \
###  lambda row: "1"

fld_funcs["meter_type"] = \
  lambda row: row['Meter_type']

fld_funcs["m_num"] = \
  lambda row: row['Meter_id']

fld_funcs["abonent_name"] = \
  lambda row: '<Название организации>"'

circuit_map = \
  {
     '220V': 'однофазний',
     '3*220V': 'прямого'
  }

tzones_map = \
  {
     '1 tar.': 'однозонний',
     '2 tar_dom': 'двохзонний',
     '2 tar. (2002)': 'двохзонний'
  }

accuracy_map = \
  {
     'GAMA 100 G1M152.220.F3': '1',
     'GAMA 100 G1M153.220.F3': '1',
     'GAMA 300 G3Y.144.230.F38': '1(2)',
     'GAMA 300 G3B.144.230.F48': '1(2)'
  }

fld_funcs["circuit"] = \
  lambda row: circuit_map[row['Circuit']]

fld_funcs["accur"] = \
  lambda row: accuracy_map[row['Meter_type']]

fld_funcs["digits"] = \
  lambda row: row['Digits'][1:].replace('d','.')


fld_funcs["tariff_zones"] = \
  lambda row: tzones_map[row['Tariff_set']]

"""
fld_funcs["tariff_zones"] = \
  lambda row: "двохзонний"
"""
def empty(str):
    return len(str.strip()) == 0

fld_funcs["seals"] = \
  lambda row: row['Seal'] if empty(row['Seal_opto']) else \
              row['Seal_opto'] if empty(row['Seal']) else \
              row['Seal'] + ', ' + row['Seal_opto']

fld_funcs["customer_name"] = \
  lambda row: ""

date_regexp = re.compile("(\d+)\.(\d+)\.(\d+)") 

def extract_date(str):
    day =  int(date_regexp.match(str).group(1))
    month =  int(date_regexp.match(str).group(2))
    year =  int(date_regexp.match(str).group(3))
    return {"day":day,"month":month,"year":year}

fld_funcs["prog_date"] = \
  lambda row: "{day:02}.{month:02}.{year:04}".format(**extract_date(row['Prog_date']))


roman = {1:"I",2:"II",3:"III",4:"IV"}
    
def quartal(month):
    return int( (month-1) / 3 ) + 1
        

def construct_issue_ver_date(issue,ver_date):
    issue_date = extract_date(issue) 
    ver_date_date = extract_date(ver_date)
    Q = roman[quartal(ver_date_date["month"])]
    return "{year1}/{Q}.{year2}".format(year1=issue_date["year"],year2=ver_date_date["year"],Q=Q) 

fld_funcs["YY/Q.YY"] = \
  lambda row: construct_issue_ver_date(row["Issue"],row["Ver_date"])

###

### readings fields ###
fld_reads = {}
fld_reads['read0'] = 0
fld_reads['read1'] = 1
fld_reads['read2'] = 2
fld_reads['read3'] = 3

### get reads_list ###
reads_list = get_data(readings_fname)
### ###

def generate_doc (tbl,data_row):
   
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:

                for fld_name, func in fld_funcs.items():
           
                    if p.text.strip() == "<"+fld_name+">":
                       p.runs[0].text = func(data_row)

                       for i in range(1,len(p.runs)):
                           p.runs[i].text = ""

def generate_readings (tbl, tblcopy, data_row):
    meter_num=data_row["Meter_id"]
    print('meter_num: ', meter_num)
    for row_idx, row in enumerate(tbl.rows):
        for cell_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for read_name, ind in fld_reads.items():
                    if p.text.strip() == "<"+read_name+">":
                        if meter_num in reads_list:
                            p.runs[0].text = reads_list[meter_num][ind]
                            row_copy = tblcopy.rows[row_idx]
                            cell_copy = row_copy.cells[cell_idx]
                            p_copy = cell_copy.paragraphs[0]
                            p_copy.runs[0].text = p.runs[0].text
                            for i in range(1,len(p_copy.runs)):
                                p_copy.runs[i].text = ""
                        else:
                            p.runs[0].text = ""
                        
meter_type_for_fname_map = \
{
'GAMA 100 G1M152.220.F3':'Gama100G1M', 
'GAMA 100 G1M153.220.F3':'Gama100G1M',
'GAMA 300 G3B.144.230.F17':'Gama300G3B',
'GAMA 300 G3B.144.230.F27':'Gama300G3B', 
'GAMA 300 G3B.144.230.F47':'Gama300G3B',
'GAMA 300 G3B.144.230.F48':'Gama300G3B',
'GAMA 300 G3B.547':'Gama300G3B',
'GAMA 300 G3M.144.230.F17':'Gama300G3M',
'GAMA 300 G3Y.144.230.F38':'Gama300G3Y'
}

tzones_for_fname_map = \
{
'1 tar.': '1tar',
'2 tar_dom': '2tar',
'2 tar. (2002)': '2tar'
}    

tpl_orig = Document (template_file_name)

with open(csv_file_name, newline='') as csv_file:
     data_rows = csv.reader (csv_file, csv_dialect)    
     fld_names = next(data_rows)
     for data_row_values in data_rows:

         data_row = dict(zip(fld_names,data_row_values))

         tpl = copy.deepcopy(tpl_orig)

         generate_doc(tpl.tables[0],data_row)

         generate_doc(tpl.tables[1],data_row)

         generate_readings(tpl.tables[0], tpl.tables[1], data_row)

         fname = "protocol_(MSO)_{meter_type}_{tzones}_{meter_num}_(p_{blanc_num}).docx"\
                 .format(\
                        meter_type=meter_type_for_fname_map[data_row["Meter_type"]],\
                        tzones=tzones_for_fname_map[data_row["Tariff_set"]],\
                        meter_num=data_row["Meter_id"],\
                        blanc_num=data_row["Doc_num"][-3:]) 
 
         tpl.save(out_folder+'/'+fname)
         

